from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_heading_style(doc, text, level):
    heading = doc.add_heading(text, level)
    return heading

def create_exam_guide():
    doc = Document()
    
    # --- Title Page ---
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run('MEDIA AND POLITICS IN PAKISTAN\n(PS 2103)')
    title_run.font.size = Pt(22)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 51, 102)
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run('\nCOMPREHENSIVE EXAM PREPARATION GUIDE\nFinal Examination Study Material')
    subtitle_run.font.size = Pt(14)
    subtitle_run.font.italic = True
    
    doc.add_paragraph('\n\nGC University Lahore\nDepartment of Political Science', style='Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()
    
    # --- Table of Contents ---
    doc.add_heading('TABLE OF CONTENTS', 1)
    toc_items = [
        'Course Overview', 'Goal 1: Introduction to Media and Politics',
        'Goal 2: Theoretical Frameworks', 'Goal 3: Media and Political Systems',
        'Goal 4: Political Communication', 'Goal 5: Print Media and Politics',
        'Goal 6: Media and Public Policy', 'Goal 7: Media Effects & Public Opinion',
        'Goal 8: Media and Elections', 'Goal 9: Media Ownership',
        'Goal 10: New Media Technologies', 'Exam Patterns & MCQ Bank'
    ]
    for item in toc_items:
        doc.add_paragraph(item, style='List Number')
    doc.add_page_break()

    # --- Goal 1 & 2: Concepts and Theory ---
    doc.add_heading('GOAL 1 & 2: FOUNDATIONS AND THEORIES', 1)
    doc.add_paragraph('Core Concepts:', style='Heading 2')
    concepts = [
        'Agenda Setting: The ability of media to influence the importance placed on topics of the public agenda.',
        'Framing: How media packages and presents information which encourages certain interpretations.',
        'Priming: Media images stimulating related thoughts in the minds of audience members.',
        'Gatekeeping: The process through which information is filtered for dissemination.'
    ]
    for concept in concepts:
        doc.add_paragraph(concept, style='List Bullet')

    # --- Goal 3 & 4: Systems and Communication ---
    doc.add_heading('GOAL 3 & 4: MEDIA SYSTEMS & COMMUNICATION', 1)
    doc.add_paragraph(
        "Pakistan operates under a complex media system that has transitioned from state-controlled "
        "to a commercialized private model (post-2002 PEMRA ordinance).", style='Normal'
    )
    doc.add_paragraph('Political Communication Strategies:', style='Heading 2')
    strategies = ['Political Branding', 'Spin Doctoring', 'Negative Campaigning', 'Grassroots Digital Mobilization']
    for s in strategies:
        doc.add_paragraph(s, style='List Bullet')

    # --- Goal 10: New Media ---
    doc.add_heading('GOAL 10: NEW MEDIA TECHNOLOGIES', 1)
    doc.add_paragraph(
        "Social media (Twitter/X, Facebook, TikTok) has bypassed traditional gatekeepers in Pakistan, "
        "allowing for rapid mobilization but also increasing the risk of 'Fake News' and polarization.", style='Normal'
    )

    # --- MCQ Bank ---
    doc.add_page_break()
    doc.add_heading('MCQ BANK', 1)
    mcqs = [
        "1. Which ordinance led to the liberalization of electronic media in Pakistan? (Ans: PEMRA Ordinance 2002)",
        "2. The 'Fourth Estate' refers to: (Ans: The Press/Media)",
        "3. Who coined the term 'Agenda Setting'? (Ans: Maxwell McCombs and Donald Shaw)",
        "4. Digital Democracy refers to: (Ans: Use of IT to enhance democratic processes)"
    ]
    for mcq in mcqs:
        doc.add_paragraph(mcq, style='Normal')

    # Save the document
    doc.save('Media_and_Politics_Exam_Guide.docx')
    print("Document created successfully: Media_and_Politics_Exam_Guide.docx")

if __name__ == "__main__":
    create_exam_guide()
